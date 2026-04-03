"""서면자문 의견서 마크다운 → DOCX 변환 스크립트."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── 경로 설정 ──
SRC = Path(__file__).parent / "서면자문_의견서_류승인_v8.md"
DST = Path(__file__).parent / "서면자문_의견서_류승인_v8.docx"

# ── 폰트 설정 ──
FONT_NAME = "맑은 고딕"
FONT_SIZE_BODY = Pt(11)
FONT_SIZE_H1 = Pt(18)
FONT_SIZE_H2 = Pt(14)
FONT_SIZE_H3 = Pt(12)
FONT_SIZE_SMALL = Pt(9)
LINE_SPACING = 1.5


def set_font(run, name=FONT_NAME, size=FONT_SIZE_BODY, bold=False, italic=False, color=None):
    """Run에 폰트 속성을 설정."""
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    # 한글 폰트 설정
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{name}"/>')
        rpr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = color


def set_paragraph_format(para, spacing_before=0, spacing_after=Pt(6), line_spacing=LINE_SPACING):
    """문단 포맷 설정."""
    pf = para.paragraph_format
    pf.space_before = Pt(spacing_before) if isinstance(spacing_before, (int, float)) else spacing_before
    pf.space_after = spacing_after
    pf.line_spacing = line_spacing


def add_styled_paragraph(doc, text, level=0, bold=False):
    """스타일이 적용된 문단 추가."""
    para = doc.add_paragraph()

    if level == 1:  # H1
        set_paragraph_format(para, spacing_before=12, spacing_after=Pt(12))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        set_font(run, size=FONT_SIZE_H1, bold=True)
    elif level == 2:  # H2
        set_paragraph_format(para, spacing_before=18, spacing_after=Pt(8))
        run = para.add_run(text)
        set_font(run, size=FONT_SIZE_H2, bold=True, color=RGBColor(0x1A, 0x56, 0xDB))
    elif level == 3:  # H3
        set_paragraph_format(para, spacing_before=12, spacing_after=Pt(6))
        run = para.add_run(text)
        set_font(run, size=FONT_SIZE_H3, bold=True)
    else:  # 본문
        set_paragraph_format(para)
        add_rich_text(para, text, bold=bold)

    return para


def add_rich_text(para, text, bold=False):
    """볼드, 이탤릭 등 인라인 마크다운 처리."""
    # **bold** 패턴 처리
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            set_font(run, bold=True)
        else:
            # `code` 패턴 처리
            code_parts = re.split(r'(`[^`]+`)', part)
            for cp in code_parts:
                if cp.startswith('`') and cp.endswith('`'):
                    run = para.add_run(cp[1:-1])
                    set_font(run, name="Consolas", size=Pt(10))
                    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
                else:
                    run = para.add_run(cp)
                    set_font(run, bold=bold)


def add_blockquote(doc, text):
    """인용 블록 추가."""
    para = doc.add_paragraph()
    set_paragraph_format(para, spacing_before=4, spacing_after=Pt(4))
    # 왼쪽 들여쓰기
    para.paragraph_format.left_indent = Cm(1.5)

    # 🔧 이모지가 있으면 그대로 유지
    clean = text.lstrip("> ").strip()
    add_rich_text(para, clean)

    # 글자색을 회색으로
    for run in para.runs:
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.font.italic = True


def add_table(doc, rows):
    """마크다운 테이블 → docx 테이블."""
    if len(rows) < 2:
        return

    # 헤더와 구분선 제거
    headers = [c.strip() for c in rows[0].strip("|").split("|")]
    data_rows = []
    for row in rows[1:]:
        if re.match(r'^\|[\s\-:]+\|', row):
            continue
        cells = [c.strip() for c in row.strip("|").split("|")]
        data_rows.append(cells)

    if not data_rows:
        return

    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더 행
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(h)
        set_font(run, size=Pt(10), bold=True)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 헤더 배경색
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F0FE"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # 데이터 행
    for r_idx, row_data in enumerate(data_rows):
        for c_idx in range(min(len(row_data), num_cols)):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            add_rich_text(para, row_data[c_idx])
            for run in para.runs:
                run.font.size = Pt(10)

    # 테이블 후 간격
    doc.add_paragraph()


def add_footnotes_section(doc, footnotes):
    """출처 섹션 추가."""
    para = doc.add_paragraph()
    set_paragraph_format(para, spacing_before=18)
    run = para.add_run("출처")
    set_font(run, size=FONT_SIZE_H2, bold=True, color=RGBColor(0x1A, 0x56, 0xDB))

    for fn in footnotes:
        para = doc.add_paragraph()
        set_paragraph_format(para, spacing_before=0, spacing_after=Pt(2))
        add_rich_text(para, fn)
        for run in para.runs:
            run.font.size = FONT_SIZE_SMALL


def convert():
    """메인 변환 로직."""
    md_text = SRC.read_text(encoding="utf-8")
    lines = md_text.split("\n")

    doc = Document()

    # ── 페이지 설정 (A4) ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── 기본 스타일 설정 ──
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE_BODY
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
        rpr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:eastAsia"), FONT_NAME)

    i = 0
    table_buffer = []
    in_table = False
    footnotes = []
    in_footnotes = False
    in_code_block = False

    while i < len(lines):
        line = lines[i]

        # 코드 블록 스킵 (```로 감싸진 부분)
        if line.strip().startswith("```"):
            if in_code_block:
                in_code_block = False
                # 코드 블록 내용은 인용 스타일로
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            # 코드 블록 내용 → 들여쓰기 + 고정폭 폰트
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1)
            set_paragraph_format(para, spacing_before=0, spacing_after=Pt(2), line_spacing=1.15)
            run = para.add_run(line)
            set_font(run, name="Consolas", size=Pt(9))
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            i += 1
            continue

        # 출처 섹션 감지
        if line.strip() == "## 출처":
            in_footnotes = True
            i += 1
            continue

        if in_footnotes:
            if line.strip():
                footnotes.append(line.strip())
            i += 1
            continue

        # 테이블 처리
        if line.strip().startswith("|"):
            table_buffer.append(line)
            i += 1
            # 다음 줄도 테이블인지 확인
            if i >= len(lines) or not lines[i].strip().startswith("|"):
                add_table(doc, table_buffer)
                table_buffer = []
            continue

        # 테이블 버퍼가 남아있으면 플러시
        if table_buffer:
            add_table(doc, table_buffer)
            table_buffer = []

        # 빈 줄
        if not line.strip():
            i += 1
            continue

        # 수평선 (---)
        if re.match(r'^-{3,}$', line.strip()):
            # 구분선 추가
            para = doc.add_paragraph()
            set_paragraph_format(para, spacing_before=6, spacing_after=Pt(6))
            # 얇은 선
            pPr = para._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # H1
        if line.startswith("# ") and not line.startswith("## "):
            add_styled_paragraph(doc, line[2:].strip(), level=1)
            i += 1
            continue

        # H2
        if line.startswith("## "):
            add_styled_paragraph(doc, line[3:].strip(), level=2)
            i += 1
            continue

        # H3
        if line.startswith("### "):
            add_styled_paragraph(doc, line[4:].strip(), level=3)
            i += 1
            continue

        # 인용 블록
        if line.startswith(">"):
            add_blockquote(doc, line)
            i += 1
            continue

        # 별첨 라인 (**[별첨...]** 형태)
        if line.strip().startswith("**[별첨"):
            para = doc.add_paragraph()
            set_paragraph_format(para, spacing_before=6)
            add_rich_text(para, line.strip())
            i += 1
            continue

        # 불릿 리스트
        if line.strip().startswith("- "):
            para = doc.add_paragraph()
            set_paragraph_format(para)
            para.paragraph_format.left_indent = Cm(1)
            para.paragraph_format.first_line_indent = Cm(-0.5)
            text = line.strip()[2:]
            run = para.add_run("• ")
            set_font(run)
            add_rich_text(para, text)
            i += 1
            continue

        # 일반 본문
        add_styled_paragraph(doc, line.strip(), level=0)
        i += 1

    # 출처 추가
    if footnotes:
        add_footnotes_section(doc, footnotes)

    doc.save(str(DST))
    print(f"변환 완료: {DST}")
    print(f"파일 크기: {DST.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    convert()
