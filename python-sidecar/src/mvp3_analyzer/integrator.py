"""MVP 3 Step 2: 통합 분석 리포트 생성.

google-genai SDK 직접 사용 (langchain 제거).
"""

import re
import json
import time
import logging
from pathlib import Path

from google import genai
from google.genai import types

from src.shared.result import Result
from src.shared.env_loader import ensure_api_key
from src.shared import read_text_file

logger = logging.getLogger("mvp")

_EVALUATOR_PROMPT = """당신은 교육시설 사전기획 리포트 **품질 검증자**입니다.
Generator가 생성한 통합 분석 리포트를 독립적으로 평가합니다.

다음 5개 항목을 점검하고, 각 항목을 PASS 또는 FAIL로 평가하십시오.
FAIL인 경우 구체적인 이유를 한 줄로 기술하십시오.

점검 항목:
1. 사업 개요 테이블: 6개 항목(사업유형/학급수/면적/사업비/기간/발주방식) 모두 기재되었는가?
2. 교육과정 운영 분석: 2-1(교과편성), 2-2(특색교육), 2-3(창의체험) 소항목이 모두 있는가?
3. 공간 기획 연계 분석: 3-1(일관성), 3-2(규모 적정성) 소항목이 모두 있는가?
4. 종합 평가: 우수사항/확인필요/검토시사점 3개 소항목이 모두 있는가?
5. 근거 인용: 분석 내용에 원문 근거 인용(괄호 표기)이 포함되어 있는가?

출력 형식 (JSON만 출력, 다른 텍스트 없음):
{"passed": true/false, "score": 0-5, "issues": ["FAIL 항목 설명", ...]}"""

SYSTEM_PROMPT = """당신은 '교육시설 사전기획 적정성 검토 전문가'입니다.

[역할]
학교 신설·증축·개축 사전기획 문서와 교육과정 운영 자료를 교차 분석하여,
교육적 타당성·공간 기획의 일관성·운영 실현 가능성을 평가합니다.

[분석 원칙]
1. 제공된 데이터만을 근거로 분석 — 추측·창작 절대 금지
2. 원문의 고유 명칭·수치를 그대로 인용
3. 데이터에 없는 항목은 '미기재'로 명시
4. 긍정적 평가와 개선 제안을 균형 있게 서술
5. 전문 용어 사용 시 괄호 안에 간략한 설명 병기"""

USER_INSTRUCTION = """아래 문서 요약 데이터를 기반으로 **통합 분석 리포트**를 작성하십시오.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
## 출력 형식 (마크다운)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

# 📋 [학교명] 사전기획 통합 분석 리포트

## 1. 사업 개요
| 항목 | 내용 |
|------|------|
| 사업유형 | (신설/증축/개축/리모델링) |
| 학급수/학생수 | (학년별 학급수, 특수학급 포함) |
| 주요면적 (부지/건축/연면적) | (㎡ 단위 기재) |
| 총사업비 | (억 원 단위 기재) |
| 사업기간 | (착공~준공 연도) |
| 설계발주방식 | (턴키/분리발주 등) |

## 2. 교육과정 운영 분석

> 이 섹션은 학교의 교육과정 편성표·교육계획서 자료를 기반으로 작성합니다.
> 자료가 없는 항목은 '해당 자료 미제공'으로 명시하고, 있는 항목은 원문을 적극 인용하십시오.

### 2-1. 교과 편성 특징
- 학년군별 교과 편성의 핵심 특징 (시수 증감, 집중이수 여부 등)
- 국가 교육과정 대비 학교 자율 운영 비율 및 특이사항

### 2-2. 특색교육 프로그램
- 학교가 중점 운영하는 특색교육 목표·내용·대상 학년
- 특색교육 실현을 위해 요구되는 공간·시설과 현재 계획과의 연계성

### 2-3. 창의적 체험활동 및 방과후
- 자율·동아리·봉사·진로 활동의 구성과 학생 참여 비율
- 방과후학교·돌봄 프로그램의 규모와 전용 공간 필요성

## 3. 공간 기획 연계 분석

### 3-1. 교육과정 ↔ 공간기획 일관성
- 특색교육 실현에 필요한 전용 공간이 스페이스 프로그램에 반영되었는지 항목별 확인
- 교육과정 요구 공간 vs. 계획 공간 정합성 표 (있을 경우)

### 3-2. 규모 적정성
- 학생 1인당 연면적, 교실당 학생수 등 교육부 기준 대비 적정성
- 특수실(과학실·도서관·체육관 등) 수량·면적의 충분성

## 4. 종합 평가

### ✅ 우수사항
- (원문 근거와 함께 3~5개, 교육과정·공간·사업 관리 측면 포함)

### ⚠️ 확인·보완 필요사항
- (원문 근거와 함께 3~5개, 항목별 개선 방향 제시)

### 💡 검토 시사점
- 사전기획 적정성 심사 시 특별히 주목할 쟁점 (법규·기준·교육부 지침 관점)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
[작성 원칙]
- 데이터에 없는 항목은 임의 추정 금지 — '해당 자료 미제공'으로 명시
- 모든 수치·명칭은 원문 그대로 인용하고 괄호로 출처 표시
- 표와 목록으로 가독성 확보, 소제목 누락 금지
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━"""


def _evaluate_report(client, model_name: str, report_text: str) -> dict:
    """생성된 리포트를 독립 컨텍스트에서 품질 검증 (Evaluator 역할).

    Generator와 분리된 새 API 호출로 편향 없는 평가를 수행합니다.
    Returns:
        {"passed": bool, "score": int, "issues": list[str]}
    """
    eval_config = types.GenerateContentConfig(
        temperature=0.0,  # 평가는 결정론적으로
        system_instruction=_EVALUATOR_PROMPT,
    )
    try:
        # 리포트 전체 구조를 평가 (섹션 4 종합평가는 후반부에 위치)
        # 최대 20,000자: Evaluator 토큰 여유 확보 + 4개 섹션 모두 커버
        _MAX_EVAL_CHARS = 20_000
        eval_text = report_text if len(report_text) <= _MAX_EVAL_CHARS else report_text[:_MAX_EVAL_CHARS]
        response = client.models.generate_content(
            model=model_name,
            contents=f"[평가 대상 리포트]\n{eval_text}",
            config=eval_config,
        )
        raw = (response.text or "").strip()
        # JSON 블록 추출: ```json ... ``` 또는 ``` ... ``` 형식 지원
        m = re.search(r"```(?:json)?\s*([\s\S]*?)```", raw)
        if m:
            raw = m.group(1).strip()
        # 코드블록 없으면 첫 '{' ~ 마지막 '}' 범위 추출
        elif "{" in raw and "}" in raw:
            raw = raw[raw.index("{") : raw.rindex("}") + 1]
        return json.loads(raw)
    except Exception as e:
        logger.warning(f"  [Evaluator] 평가 실패 (무시하고 계속): {e}")
        return {"passed": False, "score": -1, "issues": ["평가 API 호출 실패 — 품질 검증 불가"]}


def run_step2(
    processed_dir: Path,
    output_dir: Path,
    model_name: str,
    temperature: float,
    timeout: int,
    wait_time: int,
    num_variants: int,
    step1_result: Result | None = None,
    cancel_event=None,
    progress_callback=None,
    evaluate: bool = True,
) -> Result:
    """요약 파일들을 통합 분석하여 최종 리포트를 생성합니다.

    evaluate=True이면 Generator-Evaluator 패턴으로 품질 검증을 수행합니다.
    """
    result = Result()
    output_dir.mkdir(parents=True, exist_ok=True)

    api_key = ensure_api_key()
    # timeout은 SDK 기본값 사용 — http_options timeout(ms)을 짧게 잡으면
    # 응답 대기 중 끊김 발생. 에러는 아래 재시도 루프(3회)에서 처리.
    client = genai.Client(api_key=api_key)
    gen_config = types.GenerateContentConfig(
        temperature=temperature,
        system_instruction=SYSTEM_PROMPT,
    )

    # summarizer 출력은 summary_*.txt — .md도 포함
    summary_files = list(processed_dir.glob("*.txt")) + list(processed_dir.glob("*.md"))
    if not summary_files:
        logger.warning(f"요약 파일이 없습니다: {processed_dir}")
        return result
    summary_files.sort()

    if step1_result and step1_result.fail_count > 0:
        warning_msg = (
            f"주의: Step 1에서 {step1_result.fail_count}개 파일 요약 실패. "
            f"총 {step1_result.total}개 중 {step1_result.success_count}개만 분석에 포함됩니다."
        )
        logger.warning(f"  {warning_msg}")
        result.warnings.append(warning_msg)

    # 빈 파일 제외 (불필요한 API 비용 방지)
    file_contents = []
    for f in summary_files:
        text = read_text_file(f)
        if text.strip():
            file_contents.append(f"[파일: {f.name}]\n{text}")
    if not file_contents:
        logger.warning(f"유효한 요약 내용이 없습니다: {processed_dir}")
        return result
    combined = "\n\n---\n\n".join(file_contents)

    # Context 크기 체크: 한국어는 평균 1.5~2자/토큰이므로 보수적으로 2자/토큰으로 추정
    _MAX_INPUT_CHARS = 120_000  # ~60,000 토큰 추정치 — Gemini Flash 컨텍스트 여유 확보
    input_chars = len(combined)
    est_tokens = input_chars // 2  # 한국어 기준 보수적 추정
    logger.info(f"  입력 크기: {input_chars:,}자 (약 {est_tokens:,} 토큰 추정)")
    if input_chars > _MAX_INPUT_CHARS:
        logger.warning(
            f"  입력이 큽니다 ({input_chars:,}자 > {_MAX_INPUT_CHARS:,}자 권장). "
            f"앞부분 {_MAX_INPUT_CHARS:,}자로 절단합니다."
        )
        # 파일 경계에서 절단 (중간에 자르지 않음)
        truncated_parts = []
        total_len = 0
        sep = "\n\n---\n\n"
        for part in file_contents:
            if total_len + len(sep) + len(part) > _MAX_INPUT_CHARS:
                break
            truncated_parts.append(part)
            total_len += len(sep) + len(part)
        if not truncated_parts:
            # 파일 1개도 못 들어가면 첫 파일만 절단해서 사용
            truncated_parts = [file_contents[0][:_MAX_INPUT_CHARS]]
        combined = sep.join(truncated_parts)
        skipped = len(file_contents) - len(truncated_parts)
        result.warnings.append(
            f"입력 크기 초과 ({est_tokens:,} 토큰 추정): {skipped}개 파일을 제외하고 분석합니다."
        )

    logger.info(f"  {len(file_contents)}개 요약 파일 통합 분석 시작")

    for i in range(num_variants):
        if cancel_event and cancel_event.is_set():
            logger.warning("사용자에 의해 취소됨")
            break
        if progress_callback:
            progress_callback(i, num_variants, f"⏳ AI 통합분석 요청 중 (최대 {timeout}초)")

        # 재시도 로직: 503/429 transient 에러는 재시도
        text = None
        for attempt in range(3):
            if cancel_event and cancel_event.is_set():
                break
            try:
                logger.info(f"  통합분석 리포트 생성 중... (시도 {attempt + 1}/3)")
                response = client.models.generate_content(
                    model=model_name,
                    contents=f"{USER_INSTRUCTION}\n\n[데이터]:\n{combined}",
                    config=gen_config,
                )
                try:
                    text = response.text or ""
                except ValueError:
                    reason = getattr(response.candidates[0], 'finish_reason', 'UNKNOWN') if response.candidates else 'NO_CANDIDATES'
                    logger.warning(f"  안전 필터 차단 (사유: {reason})")
                    text = ""
                break  # 성공하면 재시도 루프 탈출

            except Exception as e:
                err_str = str(e).lower()
                is_transient = any(kw in err_str for kw in ("503", "429", "unavailable", "quota", "timeout"))
                if is_transient and attempt < 2:
                    retry_wait = 5 * (2 ** attempt)  # wait_time(variant 간 대기)과 구분
                    logger.warning(f"  임시 오류 (재시도 {retry_wait}초 후): {e}")
                    if progress_callback:
                        progress_callback(i, num_variants, f"⏳ API 대기 중... {retry_wait}초")
                    if cancel_event:
                        if cancel_event.wait(retry_wait):
                            break
                    else:
                        time.sleep(retry_wait)
                else:
                    logger.error(f"  생성 실패: {e}")
                    result.add_failure(f"분석 결과 {i + 1}", str(e))
                    if progress_callback:
                        progress_callback(i, num_variants, f"✕ 분석 실패: {str(e)[:100]}")
                    break

        if text is None:
            continue

        if not text:
            logger.warning(f"  빈 응답 (안전 필터 또는 빈 결과)")
            result.add_failure(f"분석 결과 {i + 1}", "AI 응답이 비어 있습니다 (안전 필터 차단 가능성)")
            continue

        if num_variants == 1:
            save_name = "통합 분석 리포트.md"
        else:
            save_name = f"통합 분석 리포트 {i + 1}.md"

        try:
            save_path = output_dir / save_name
            save_path.write_text(text, encoding="utf-8")

            # Evaluator: Generator와 독립 컨텍스트에서 품질 검증
            if evaluate and not (cancel_event and cancel_event.is_set()):
                if progress_callback:
                    progress_callback(i + 1, num_variants, "🔍 품질 검증 중...")
                eval_result = _evaluate_report(client, model_name, text)
                score = eval_result.get("score", -1)
                issues = eval_result.get("issues", [])
                passed = eval_result.get("passed", True)
                if score >= 0:
                    logger.info(f"  [Evaluator] 점수: {score}/5 — {'PASS' if passed else 'FAIL'}")
                if not passed and issues:
                    warn_msg = f"리포트 품질 검증 미흡 ({score}/5점): " + "; ".join(issues)
                    logger.warning(f"  [Evaluator] {warn_msg}")
                    result.warnings.append(warn_msg)
                elif score >= 0:
                    logger.info(f"  [Evaluator] 품질 검증 통과")

            result.add_success(save_name, str(save_path))
            if progress_callback:
                progress_callback(i + 1, num_variants, f"✓ 통합분석 완료")
        except Exception as e:
            logger.error(f"  파일 저장 실패: {e}")
            result.add_failure(f"분석 결과 {i + 1}", f"파일 저장 실패: {e}")

        if i < num_variants - 1:
            if cancel_event:
                if cancel_event.wait(wait_time):
                    break
            else:
                time.sleep(wait_time)

    # 산출물 폴더 정리: 리포트를 최상단으로, 중간산물은 _작업데이터로
    _reorganize_output_structure(output_dir, result)

    # 요약 시트 (구조화된 Excel + MD 백업)
    _export_summary_sheet(summary_files, output_dir, result)

    # Markdown → DOCX 변환
    _convert_reports_to_docx(output_dir, result)

    if progress_callback:
        progress_callback(num_variants, num_variants, "통합분석 완료")
    result.output_path = str(output_dir)
    return result


def _export_summary_sheet(summary_files: list, output_dir: Path, result) -> None:
    """Step1 요약 파일들을 구조화된 Excel + MD 백업으로 내보냅니다."""
    if not summary_files:
        return
    try:
        # MD 백업은 _작업데이터/에 저장
        parts = []
        for f in sorted(summary_files):
            text = read_text_file(f)
            if text.strip():
                origin = f.stem.removeprefix("summary_")
                parts.append(f"## {origin}\n\n{text.strip()}")
        if not parts:
            return
        work_dir = output_dir / "_작업데이터"
        work_dir.mkdir(parents=True, exist_ok=True)
        sheet_path = work_dir / "사전기획 요약 시트.md"
        sheet_path.write_text("\n\n---\n\n".join(parts), encoding="utf-8")

        # 구조화된 Excel은 최상단에 저장
        _export_summary_sheet_xlsx(summary_files, output_dir, result)
    except Exception as e:
        logger.warning(f"  요약 시트 저장 실패 (무시): {e}")


def _parse_summary_sections(text: str) -> list[dict]:
    """마크다운 요약 텍스트를 섹션별 키-값 쌍으로 파싱합니다.

    Returns:
        [{"섹션": "기본정보", "항목": "부지면적", "내용": "25,273.00㎡"}, ...]
    """
    rows = []
    current_section = ""

    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue

        # 섹션 헤더 감지: "## N. 제목" 또는 "**N. 제목**" 또는 "# 제목"
        section_match = re.match(r'^#{1,3}\s*\d*\.?\s*(.+?)$', line)
        if section_match:
            current_section = section_match.group(1).strip().rstrip("*").strip()
            continue
        # "**섹션명**" 단독 행도 섹션 헤더로 인식
        bold_section = re.match(r'^\*\*(\d+\.?\s*.+?)\*\*\s*$', line)
        if bold_section and ":" not in line and "：" not in line:
            current_section = bold_section.group(1).strip()
            continue

        # 키-값 쌍 감지: "**키**: 값" 또는 "* **키**: 값" 또는 "- 키: 값"
        kv_match = re.match(
            r'^[\*\-·•]?\s*\*?\*?([^*:：]+?)\*?\*?\s*[:：]\s*(.+)$', line
        )
        if kv_match:
            key = kv_match.group(1).strip().lstrip("*").rstrip("*").strip()
            val = kv_match.group(2).strip()
            if key and val and len(key) < 40:
                rows.append({
                    "섹션": current_section or "기타",
                    "항목": key,
                    "내용": val,
                })
                continue

        # 목록 항목: "* 내용" 또는 "- 내용" (키-값이 아닌 것)
        list_match = re.match(r'^[\*\-·•]\s+(.+)$', line)
        if list_match and current_section:
            content = list_match.group(1).strip()
            if content and len(content) > 3:
                rows.append({
                    "섹션": current_section,
                    "항목": "",
                    "내용": content,
                })

    return rows


def _export_summary_sheet_xlsx(summary_files: list, output_dir: Path, result) -> None:
    """Step1 요약 파일들을 구조화된 Excel로 내보냅니다.

    시트 구성:
    - 종합: 파일별 주요 항목을 행-열로 정리
    - 파일별 시트: 섹션/항목/내용 구조화
    """
    if not summary_files:
        return
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter

        wb = Workbook()

        # ── 스타일 정의 ──
        BLUE = "1F4E79"
        LIGHT_BLUE = "D6E4F0"
        WHITE = "FFFFFF"
        GRAY_BG = "F2F2F2"

        header_fill = PatternFill(start_color=BLUE, end_color=BLUE, fill_type="solid")
        header_font = Font(bold=True, color=WHITE, size=10)
        section_fill = PatternFill(start_color=LIGHT_BLUE, end_color=LIGHT_BLUE, fill_type="solid")
        section_font = Font(bold=True, size=10)
        alt_fill = PatternFill(start_color=GRAY_BG, end_color=GRAY_BG, fill_type="solid")
        cell_align = Alignment(horizontal="left", vertical="top", wrap_text=True)
        center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
        border = Border(
            left=Side(style="thin", color="BFBFBF"),
            right=Side(style="thin", color="BFBFBF"),
            top=Side(style="thin", color="BFBFBF"),
            bottom=Side(style="thin", color="BFBFBF"),
        )

        def style_header(ws, row_num, max_col):
            for c in range(1, max_col + 1):
                cell = ws.cell(row=row_num, column=c)
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = center_align
                cell.border = border

        def style_cell(ws, row_num, col, is_alt=False, is_section=False):
            cell = ws.cell(row=row_num, column=col)
            cell.alignment = cell_align
            cell.border = border
            cell.font = Font(size=10)
            if is_section:
                cell.fill = section_fill
                cell.font = section_font
            elif is_alt:
                cell.fill = alt_fill

        # ── 시트 1: 종합 (파일별 핵심 항목 비교표) ──
        ws_main = wb.active
        ws_main.title = "종합"

        # 핵심 추출 항목 (요약에서 자주 등장하는 키)
        KEY_FIELDS = [
            "사업 유형", "사업유형",
            "학급수", "학생수", "학급수/학생수", "학생수 배치계획",
            "부지면적", "건축면적", "연면적",
            "총사업비", "총사업비 내역", "사업비",
            "사업 기간", "사업기간",
            "설계발주방식", "설계 발주방식",
            "교육과정 비전", "핵심 특색 프로그램",
        ]

        # 파일별 파싱
        file_data = []
        for f in sorted(summary_files):
            text = read_text_file(f)
            if not text.strip():
                continue
            origin = f.stem.removeprefix("summary_")
            parsed = _parse_summary_sections(text)
            file_data.append((origin, text, parsed))

        if not file_data:
            wb.close()
            return

        # 종합 시트: 행=항목, 열=파일
        ws_main["A1"] = "항목"
        for ci, (origin, _, _) in enumerate(file_data, start=2):
            ws_main.cell(row=1, column=ci, value=origin)
        style_header(ws_main, 1, len(file_data) + 1)

        # 각 핵심 항목에 대해 파일별 값 채우기
        used_fields = []
        for field in KEY_FIELDS:
            values = []
            for _, _, parsed in file_data:
                val = ""
                for row in parsed:
                    item = row["항목"]
                    if item and (field in item or item in field):
                        val = row["내용"]
                        break
                values.append(val)
            if any(values):  # 하나라도 값이 있으면 포함
                used_fields.append((field, values))

        for ri, (field, values) in enumerate(used_fields, start=2):
            ws_main.cell(row=ri, column=1, value=field)
            style_cell(ws_main, ri, 1, is_alt=(ri % 2 == 0))
            for ci, val in enumerate(values, start=2):
                ws_main.cell(row=ri, column=ci, value=val)
                style_cell(ws_main, ri, ci, is_alt=(ri % 2 == 0))

        # 열 너비
        ws_main.column_dimensions["A"].width = 22
        for ci in range(2, len(file_data) + 2):
            ws_main.column_dimensions[get_column_letter(ci)].width = 40
        ws_main.freeze_panes = "B2"

        # ── 파일별 시트: 섹션/항목/내용 구조화 ──
        for origin, text, parsed in file_data:
            # 시트명은 31자 제한 + 특수문자 제거
            sheet_name = re.sub(r'[\\/*?\[\]:]', '', origin)[:31]
            ws = wb.create_sheet(title=sheet_name)

            ws["A1"] = "섹션"
            ws["B1"] = "항목"
            ws["C1"] = "내용"
            style_header(ws, 1, 3)

            if parsed:
                prev_section = ""
                row_num = 2
                for item in parsed:
                    section = item["섹션"]
                    is_new_section = section != prev_section
                    ws.cell(row=row_num, column=1, value=section if is_new_section else "")
                    ws.cell(row=row_num, column=2, value=item["항목"])
                    ws.cell(row=row_num, column=3, value=item["내용"])
                    for c in range(1, 4):
                        style_cell(ws, row_num, c,
                                   is_alt=(row_num % 2 == 0),
                                   is_section=(is_new_section and c == 1))
                    prev_section = section
                    row_num += 1
            else:
                # 파싱 실패 시 원문 그대로 (폴백)
                ws.cell(row=2, column=1, value="전체")
                ws.cell(row=2, column=2, value="요약 원문")
                ws.cell(row=2, column=3, value=text.strip()[:32000])
                for c in range(1, 4):
                    style_cell(ws, 2, c)

            ws.column_dimensions["A"].width = 20
            ws.column_dimensions["B"].width = 25
            ws.column_dimensions["C"].width = 60
            ws.freeze_panes = "A2"

        # 최상단에 저장
        xlsx_path = output_dir / "사전기획 요약 시트.xlsx"
        wb.save(str(xlsx_path))

        result.add_success("사전기획 요약 시트.xlsx", str(xlsx_path))
        logger.info(f"  요약 시트(Excel) 저장 → {xlsx_path.name}")
    except Exception as e:
        logger.warning(f"  요약 시트 Excel 저장 실패 (무시): {e}")


def _reorganize_output_structure(output_dir: Path, result) -> None:
    """산출물을 사용자 관점에서 정리합니다.

    새 구조:
        output_dir/
        ├── 통합 분석 리포트.md          ← 최상단 (가장 중요)
        ├── 통합 분석 리포트.docx
        ├── 사전기획 요약 시트.xlsx
        ├── 테마별_분리_자료/            ← 테마별 PDF
        └── _작업데이터/                 ← 중간산물 (숨김)
            ├── pipeline_state.json
            ├── 사전기획 요약 시트.md
            └── 요약_중간결과/
    """
    if not output_dir.exists():
        return

    try:
        work_dir = output_dir / "_작업데이터"
        work_dir.mkdir(parents=True, exist_ok=True)
        summary_cache_dir = work_dir / "요약_중간결과"
        summary_cache_dir.mkdir(parents=True, exist_ok=True)

        # 1. 통합 분석 리포트는 최상단에 유지 (이미 output_dir에 저장됨)
        #    기존 최종_산출물/ 안에 있으면 꺼내기
        old_final = output_dir / "최종_산출물"
        if old_final.exists():
            for f in list(old_final.iterdir()):
                dest = output_dir / f.name
                if not dest.exists():
                    f.replace(dest)
            # 빈 폴더 제거
            try:
                old_final.rmdir()
            except OSError:
                pass

        # 2. 중간산물 → _작업데이터/ 이동
        #    summary_*.txt/md, pipeline_state.json, OCR 텍스트 등
        for f in list(output_dir.iterdir()):
            if f.is_dir():
                continue
            name = f.name
            # 최상단에 남길 파일: 리포트, Excel
            if "통합" in name and "리포트" in name:
                continue
            if name.endswith(".xlsx"):
                continue
            # 나머지 → _작업데이터/
            if name.startswith("summary_") or name == "pipeline_state.json":
                dest = summary_cache_dir / name if name.startswith("summary_") else work_dir / name
                f.replace(dest)
            elif name.endswith((".txt", ".md")) and "통합" not in name and "리포트" not in name:
                # OCR 결과 텍스트, 요약 시트 MD 등
                dest = work_dir / name
                f.replace(dest)

        # 3. 기존 processed/, output/ 폴더 정리 (하위호환)
        for old_name in ("processed", "output"):
            old_dir = output_dir / old_name
            if old_dir.exists() and old_dir.is_dir():
                for f in list(old_dir.iterdir()):
                    dest = summary_cache_dir / f.name if f.name.startswith("summary_") else work_dir / f.name
                    if not dest.exists():
                        f.replace(dest)
                try:
                    old_dir.rmdir()
                except OSError:
                    pass

        logger.info("  산출물 폴더 정리 완료 (리포트 최상단, 중간산물 → _작업데이터/)")

    except Exception as e:
        logger.warning(f"  폴더 정리 실패 (무시): {e}")


def _convert_reports_to_docx(output_dir: Path, result) -> None:
    """Markdown 리포트를 DOCX로 변환합니다 (python-docx 사용, 외부 의존 없음)."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.info("  python-docx 미설치 — DOCX 변환 스킵")
        return

    for md_file in sorted(output_dir.glob("통합 분석 리포트*.md")):
        try:
            md_text = md_file.read_text(encoding="utf-8")
            docx_file = md_file.with_suffix(".docx")

            doc = Document()

            # 페이지 설정: A4, 적절한 여백
            section = doc.sections[0]
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

            # 기본 스타일 설정
            style = doc.styles["Normal"]
            style.font.name = "맑은 고딕"
            style.font.size = Pt(10)
            style.paragraph_format.space_after = Pt(4)
            style.paragraph_format.line_spacing = 1.3

            # 테이블 상태
            table_rows: list[list[str]] = []
            in_table = False

            def flush_table():
                nonlocal table_rows, in_table
                if not table_rows:
                    return
                # 구분선 행(|---|---|) 제거
                data_rows = [r for r in table_rows if not all(
                    c.strip().replace("-", "").replace(":", "") == "" for c in r
                )]
                if len(data_rows) < 2:
                    table_rows = []
                    in_table = False
                    return
                cols = len(data_rows[0])
                tbl = doc.add_table(rows=len(data_rows), cols=cols)
                tbl.style = "Table Grid"
                for ri, row_data in enumerate(data_rows):
                    for ci, cell_text in enumerate(row_data):
                        if ci < cols:
                            cell = tbl.cell(ri, ci)
                            cell.text = cell_text.strip()
                            cell.paragraphs[0].paragraph_format.space_after = Pt(1)
                            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                            if run:
                                run.font.size = Pt(9)
                                if ri == 0:  # 헤더 행
                                    run.bold = True
                # 헤더 행 배경색
                for ci in range(cols):
                    shading = tbl.cell(0, ci)._element
                    from docx.oxml.ns import qn
                    tc_pr = shading.get_or_add_tcPr()
                    shd = tc_pr.makeelement(qn("w:shd"), {
                        qn("w:val"): "clear",
                        qn("w:fill"): "1F4E79",
                    })
                    tc_pr.append(shd)
                    # 헤더 텍스트 흰색
                    for run in tbl.cell(0, ci).paragraphs[0].runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

                table_rows = []
                in_table = False

            for line in md_text.split("\n"):
                stripped = line.strip()

                # 테이블 행 감지
                if stripped.startswith("|") and stripped.endswith("|"):
                    cells = [c.strip() for c in stripped.split("|")[1:-1]]
                    table_rows.append(cells)
                    in_table = True
                    continue
                elif in_table:
                    flush_table()

                # 빈 줄
                if not stripped:
                    continue

                # 제목
                if stripped.startswith("# "):
                    flush_table()
                    p = doc.add_heading(stripped[2:].strip(), level=1)
                    p.runs[0].font.size = Pt(16) if p.runs else None
                elif stripped.startswith("## "):
                    flush_table()
                    p = doc.add_heading(stripped[3:].strip(), level=2)
                    if p.runs:
                        p.runs[0].font.size = Pt(13)
                elif stripped.startswith("### "):
                    flush_table()
                    p = doc.add_heading(stripped[4:].strip(), level=3)
                    if p.runs:
                        p.runs[0].font.size = Pt(11)

                # 인용문
                elif stripped.startswith("> "):
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(1)
                    run = p.add_run(stripped[2:])
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

                # 목록
                elif stripped.startswith("- ") or stripped.startswith("* "):
                    text = stripped[2:]
                    # 볼드 처리: **텍스트**
                    p = doc.add_paragraph(style="List Bullet")
                    _add_md_runs(p, text)

                # 구분선
                elif stripped in ("---", "***", "___"):
                    doc.add_paragraph("─" * 50).runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

                # 일반 텍스트
                else:
                    p = doc.add_paragraph()
                    _add_md_runs(p, stripped)

            flush_table()

            doc.save(str(docx_file))
            result.add_success(docx_file.name, str(docx_file))
            logger.info(f"  Markdown → DOCX: {md_file.name} → {docx_file.name}")
        except Exception as e:
            logger.warning(f"  DOCX 변환 실패: {md_file.name} — {e}")


def _add_md_runs(paragraph, text: str):
    """마크다운 인라인 서식(**볼드**, *이탤릭*)을 docx Run으로 변환합니다."""
    import re as _re
    # **bold** 패턴 분리
    parts = _re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # *italic* 패턴
            sub_parts = _re.split(r'(\*[^*]+\*)', part)
            for sp in sub_parts:
                if sp.startswith("*") and sp.endswith("*") and len(sp) > 2:
                    run = paragraph.add_run(sp[1:-1])
                    run.italic = True
                elif sp:
                    paragraph.add_run(sp)
